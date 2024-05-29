import json
import boto3
from docx import Document

s3 = boto3.client('s3')
kendra = boto3.client('kendra')

def extract_tables_from_docx(docx_path):
    doc = Document(docx_path)
    tables_data = []
    for table in doc.tables:
        table_data = []
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                row_data[headers[idx]] = cell.text.strip()
            table_data.append(row_data)
        tables_data.append(table_data)
    return tables_data

def convert_tables_to_json(tables_data):
    json_data = []
    for table_data in tables_data:
        json_data.append(json.dumps(table_data, indent=4))
    return json_data

def replace_tables_with_json(docx_path, json_text, output_docx_path):
    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    json_index = 0
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "":
            if json_index < len(json_text):
                paragraph.text = json_text[json_index]
                json_index += 1
            else:
                break
    doc.save(output_docx_path)

def lambda_handler(event, context):
    # Get the bucket name and file key from the S3 event
    bucket_name = event['Records'][0]['s3']['bucket']['name']
    file_key = event['Records'][0]['s3']['object']['key']
    
    # Download the file from S3
    download_path = '/tmp/{}'.format(file_key)
    s3.download_file(bucket_name, file_key, download_path)
    
    # Process the file
    tables_data = extract_tables_from_docx(download_path)
    json_text = convert_tables_to_json(tables_data)
    output_file_key = 'modified_{}'.format(file_key)
    output_path = '/tmp/{}'.format(output_file_key)
    replace_tables_with_json(download_path, json_text, output_path)
    
    # Upload the modified file to S3
    modified_bucket_name = 'modified-word-doc-bucket'
    s3.upload_file(output_path, modified_bucket_name, output_file_key)
    
    # Ingest the modified document into Amazon Kendra. Replace index_id and role_arn with your actual value.
    index_id = 'myKendraIndexID'
    role_arn = 'myKendraRole'
    document_path = f's3://{modified_bucket_name}/{output_file_key}'
    kendra.batch_put_document(
        IndexId=index_id,
        RoleArn=role_arn,
        Documents=[
            {
                'Id': output_file_key,
                'S3Path': {
                    'Bucket': modified_bucket_name,
                    'Key': output_file_key
                }
            }
        ]
    )
    
    return {
        'statusCode': 200,
        'body': json.dumps('File processed, uploaded to S3, and ingested into Kendra: {}'.format(output_file_key))
    }
